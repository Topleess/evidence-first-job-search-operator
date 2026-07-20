#!/usr/bin/env python3
"""Render a validated tailored package as a single-column ATS-safe DOCX."""
from __future__ import annotations

from pathlib import Path
from typing import Any
from zipfile import ZipFile


class RenderValidationError(RuntimeError):
    pass


def render_docx(package: dict[str, Any], profile: dict[str, Any], output: str | Path) -> Path:
    from docx import Document
    from docx.shared import Pt

    output = Path(output)
    output.parent.mkdir(parents=True, exist_ok=True)
    candidate = profile["candidate"]
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    document.add_heading(candidate["public_name_ru"], level=0)
    document.add_paragraph(package["resume"]["target_title"])
    links = candidate.get("links") or {}
    contact_line = " | ".join(
        str(value) for value in (
            links.get("telegram_public_handle"), links.get("linkedin"), links.get("portfolio_primary")
        ) if value
    )
    if contact_line:
        document.add_paragraph(contact_line)
    document.add_heading("Профессиональный профиль", level=1)
    document.add_paragraph(package["resume"]["summary"])
    document.add_heading("Ключевые результаты", level=1)
    for bullet in package["resume"].get("evidence_bullets") or []:
        document.add_paragraph(bullet, style="List Bullet")
    document.add_heading("Опыт работы", level=1)
    for item in profile.get("experience_highlights") or []:
        document.add_heading(f"{item['role']} — {item['company']}", level=2)
        document.add_paragraph(item["period"])
        for bullet in item.get("highlights_ru") or []:
            document.add_paragraph(str(bullet), style="List Bullet")
    document.add_heading("Навыки, релевантные вакансии", level=1)
    document.add_paragraph(" • ".join(package["resume"].get("ats_keywords") or []))
    education = profile.get("education") or {}
    document.add_heading("Образование", level=1)
    document.add_paragraph("Неоконченное высшее образование")
    for item in education.get("items") or []:
        document.add_paragraph(f"{item.get('institution', '')} — {item.get('field', '')}")
    document.add_heading("Языки", level=1)
    document.add_paragraph(f"Английский — {candidate.get('english_level', 'не указан')}")
    document.save(output)
    validate_docx(output, package)
    return output


def extract_docx_text(path: str | Path) -> str:
    from docx import Document
    return "\n".join(p.text for p in Document(str(path)).paragraphs if p.text.strip())


def validate_docx(path: str | Path, package: dict[str, Any]) -> None:
    path = Path(path)
    with ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    forbidden = {"table": "<w:tbl", "text box": "<w:txbxContent", "drawing": "<w:drawing"}
    found = [name for name, marker in forbidden.items() if marker in xml]
    if found:
        raise RenderValidationError(f"ATS-unsafe DOCX elements: {', '.join(found)}")
    text = extract_docx_text(path)
    required = [
        package["resume"]["target_title"], package["resume"]["summary"],
        *(package["resume"].get("evidence_bullets") or []),
        *(package["resume"].get("ats_keywords") or []),
    ]
    missing = [item for item in required if item and item not in text]
    if missing:
        raise RenderValidationError(f"DOCX text extraction lost {len(missing)} required items")
